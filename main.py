from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from loguru import logger

from str_util import find

'''
进行以下检测

段落
    1. 首行缩进两字符
    2. 段间距相同
    3. 两端对齐
    4. 英文标点
    5. 字体 宋体，Times Roman


引用
    5. 引用连续
    6. 格式正确


············
nlp
句法错误。

图  检查居中和缩进

'''


#
# for i,p in enumerate(document.paragraphs):
#
#         print(str(i)+p.text)
#


# 检测首行缩进
def check_start(document: Document):
    res = []
    ps = get_main_p(document)
    for p in ps:

        indent = p.paragraph_format.first_line_indent
        if p.text and len(p.text) > 80:  # 太短的话应该不是一段，忽略掉
            if str.isnumeric(p.text[1]):
                continue
            if p.text[0:2] in ["注：", "Ke", "KE"]:  # 注释,关键字不用缩进
                continue
            if " <w:numPr>" in p._element.xml:  # 这说明是段落起始是标号，一般不用缩进
                if not indent:
                    continue
            if not indent in [304800, 266700, 306070, None]:  # 首行缩进要么是0.74厘米要么是两个字符
                res.append(" \"" + p.text[0:20] + "\" " + " 这一段 缩进不太对")
    return res


# 检查主题段落是否为两端对齐
def check_alignment(document: Document):
    res = []
    ps = document.paragraphs
    for p in ps:
        if p.text and len(p.text) > 80:  # 太短的话应该不是一段，忽略掉

            if p.alignment not in [WD_PARAGRAPH_ALIGNMENT.JUSTIFY, None]:  # 如果是None表示用了默认格式
                res.append(" 这一段没用两端对齐 " + p.text[0:20])
    return res


# 检查行间距
def check_line_space(document: Document):
    ps = document.paragraphs
    res = []
    for p in ps:
        if p.text and len(p.text) > 80:  # 太短的话应该不是一段，忽略掉
            if "本人声明" in p.text or "东南大学" in p.text:  # 模版里边的数据
                continue
            if p.paragraph_format.line_spacing not in [1.25, 1.5, None]:
                res.append("行间距不太对" + " \"" + p.text[0:20] + "\" ")
    return res


# 检查标点符号
def check_punctuation(documennt):
    ps = documennt.paragraphs
    paragraphs = []
    for p in ps:
        if p.text and len(p.text) > 10:  # 太短的话应该不是一段，忽略掉
            if len(p.text) < 80:
                if p.text[0] in ["图", "表", "第", "(", "[", "（"]:
                    continue
                if str.isnumeric(p.text[0]):
                    continue
        paragraphs.append(p)

    en_fenhao = ';'
    en_juhao = '.'
    en_maohao = ':'
    res = []

    def is_special(char):
        english_check = re.compile(r'[a-zA-Z0-9]')
        if english_check.match(char):
            return True
        if char in ['[', ']', ' ']:
            return True
        return False

    for p in paragraphs:

        letters = len(re.sub("[^a-zA-Z]", "", p.text))
        if letters * 2 > len(p.text):  # 英文摘要或者引用
            continue

        if en_fenhao in p.text:
            res.append("正文一般不会有英文分号吧" + " \"" + p.text[0:20] + "\" ")
        if en_maohao in p.text:
            res.append("正文一般不会有英文冒号吧" + " \"" + p.text[0:20] + "\" ")
        if ',' in p.text:
            indexs = find(p.text, ',')
            for idx in indexs:
                if idx != len(p.text) - 1:
                    if not is_special(p.text[idx + 1]):
                        res.append(" 这里不该有英文逗号吧 " + " \"" + p.text[0:20] + "\" " + str(idx))
                        continue
                if idx != 0:
                    if not is_special(p.text[idx - 1]):
                        res.append(" 这里不该有英文逗号吧 " + " \"" + p.text[0:20] + "\" " + str(idx))
                        continue

        if en_juhao in p.text:
            indexs = find(p.text, en_juhao)
            for idx in indexs:
                if idx != len(p.text) - 1:
                    if not is_special(p.text[idx + 1]):
                        res.append(" 这里不该有英文句号吧 " + " \"" + p.text[0:20] + "\" " + str(idx))
                        break
                if idx != 0:
                    if not is_special(p.text[idx - 1]):
                        res.append(" 这里不该有英文句号吧 " + " \"" + p.text[0:20] + "\" " + str(idx))
                        break
    return res


def get_main_p(document, min_chars=20):
    ps = document.paragraphs
    res = []
    for p in ps:
        if p.text and len(p.text) > min_chars:  # 太短的话应该不是一段，忽略掉
            if len(p.text) < 150:
                if p.text[0] in ["图", "表", "第", "(", "[", "（", "K", "东", "论", "日", "关", "研", "导"]:  # 这些段落不做检测
                    continue
                if str.isnumeric(p.text[0]):
                    continue

            res.append(p)
    return res


# 这里不一定能检测出来。可能要深入xml来检测
def chck_font(document):
    res = []
    font_set = {None, 'Times New Roman', '宋体', 'Arial Unicode MS'}
    paragraphs = get_main_p(document)
    for p in paragraphs:
        if p.style.font.name not in font_set:
            res.append("字体不太对吧 " + p.style.font.name + " \"" + p.text[0:20] + "\" ")

    return res


# 检查段尾是否为句号
def check_paragraph_tail(document):
    res = []
    paragraphs = get_main_p(document, min_chars=80)
    for p in paragraphs:
        if p.text[-1] not in ['。', "：", "；", ".", ":"]:
            if p.text[-1] == " ":
                res.append("段尾是空格，删掉比较好。 " + " \"" + p.text[0:20] + "\" ")
            else:
                res.append("段尾不是句号、冒号或分号，值得警惕 " + " \"" + p.text[0:20] + "\" ")
    return res


import re


def check_ref(document):
    res_arr = []
    current_max = 0
    ps = document.paragraphs
    need_start = False
    for p in ps:
        if p.text and "未找到引用源" in p.text:
            res_arr.append("这一段未找到引用源 " + p.text[0:10])

        if (p.text == "表格目录"):
            need_start = True
            continue
        elif not need_start:
            continue
        ress = re.findall("\[[0-9]+\]", p.text)
        for res in ress:
            idx = int(res[1:-1])
            if idx > current_max:
                if idx - current_max == 1:
                    current_max += 1
                else:
                    res_arr.append("引用不连续 " + " \"" + p.text[0:20] + "\" ")
    return res_arr


def check_pic_align(document):
    res = []
    ps = document.paragraphs
    last_not_center = False
    last_indent = False
    for p in ps:
        if "<w:drawing>" in p._element.xml:
            if "<w:jc w:val=\"center\"/>" not in p._element.xml:
                last_not_center = True
            if "<w:ind" in p._element.xml:
                if "w:firstLine=\"0\"" not in p._element.xml \
                        and "w:firstLineChars=\"0\"" not in p._element.xml:
                    last_indent = True
        elif p.text:
            if last_not_center:
                last_not_center = False
                res.append(p.text + " 上边的图片没居中")
            if last_indent:
                last_indent = False
                res.append(p.text + " 上边的图片不该有缩进")

    return res


def check_pic_table_desc_space(document):
    res = []
    ps = document.paragraphs
    for p in ps:
        if p.text and len(p.text) < 60:

            if len(str.strip(p.text)) > 0:
                if str.strip(p.text)[0] in ["图", "表"]:
                    if '\t' in p.text and (str.isnumeric(p.text[-1])):  # 说明是目录
                        continue
                    if str.strip(p.text)[1] in [" ", " "]:
                        res.append(p.text + " 题注里不该有空格")

                    if "<w:jc w:val=\"center\"/>" not in p._element.xml:
                        if "<w:jc w:val=\"left\"/>" in p._element.xml or "<w:jc w:val=\"right\"/>" in p._element.xml:
                            res.append(p.text + " 题注没居中")
                    if "<w:ind" in p._element.xml:
                        if "w:firstLine=\"0\"" not in p._element.xml \
                                and "w:firstLineChars=\"0\"" not in p._element.xml:
                            res.append(p.text + " 题注不该有缩进")

    return res


def check_all(doc_name):
    res = []

    cf = [check_line_space, check_start, check_alignment,
          check_punctuation, check_ref, check_paragraph_tail, check_pic_align
        , check_pic_table_desc_space]
    for c in cf:
        res_part = check(c, doc_name)
        if res_part and len(res_part) != 0:
            res.append(res_part)
    return res


def check(check_fun, doc_name):
    document = Document(doc_name)
    res_par = check_fun(document)
    return res_par

check_pic_table_desc_space(Document("data/zyk.docx"))